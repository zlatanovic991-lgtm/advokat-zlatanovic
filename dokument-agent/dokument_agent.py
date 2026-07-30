#!/usr/bin/env python3
"""
Dokument agent - automatsko sortiranje dokumenata sa Desktopa u foldere klijenata

Sta radi:
  1. Skenira Desktop (samo fajlove, ne foldere) i trazi nove dokumente
     (.pdf, .docx, .doc, .jpg, .jpeg, .png).
  2. Iz svakog dokumenta izvlaci tekst (PDF/Word tekst, ili OCR za skenirane
     PDF-ove i slike).
  3. Salje tekst lokalnom Ollama modelu koji prepoznaje tip dokumenta i
     ime i prezime klijenta na koga se dokument odnosi.
  4. Ako prepoznati klijent (fazi) odgovara postojecem folderu - premesta
     dokument tamo. Ako ne postoji nijedan slican folder - kreira novi
     folder "Prezime Ime" i premesta dokument tamo. Novo ime fajla:
     "Tip dokumenta - Prezime Ime - DD.MM.YYYY.ext".
  5. Ako je AI nesiguran (dvosmisleno ime, nejasan tip, tekst se ne moze
     procitati, ili je prepoznato ime "na granici" slicnosti sa postojecim
     folderom) - dokument se NE premesta automatski u folder klijenta, vec
     u "_Za proveru" folder na Desktopu, uz .txt napomenu zasto je sporno.

Upotreba:
  python dokument_agent.py
  python dokument_agent.py --model qwen3.6
  python dokument_agent.py --desktop "C:\\Users\\Vladimir\\OneDrive\\Desktop"

Zahteva pokrenut Ollama (ollama serve) i preuzet model (ollama pull <model>).
Za OCR: instaliran Tesseract OCR (sa srpskim jezickim paketom).
"""

import argparse
import json
import os
import shutil
import sys
import unicodedata
import urllib.error
import urllib.request
from datetime import datetime
from difflib import SequenceMatcher

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DESKTOP_DIR_DEFAULT = r"C:\Users\Vladimir\OneDrive\Desktop"
KLIJENTI_FOLDER_NAZIV = "Advokat Zlatanovic"
SPORNO_FOLDER_NAZIV = "_Za proveru"

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL_DEFAULT = "qwen3.6"

DOZVOLJENE_EKSTENZIJE = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}
MIN_STAROST_SEK = 30       # ne diraj fajl mladji od ovoga (moguce jos u toku snimanja)
MAX_TEKST_ZNAKOVA = 8000   # koliko teksta se salje modelu
MIN_TEKST_ZNAKOVA = 20     # ispod ovoga - tekst je prazan/nečitak

PODUDARANJE_POSTOJECI = 0.97   # >= ovoga: isti klijent, koristi postojeci folder
PODUDARANJE_SLICNO    = 0.75   # >= ovoga (a < gornjeg): sumnjivo slicno - u proveru

ZABRANJENI_ZNACI_FAJL = '<>:"/\\|?*'

SISTEM_PROMPT = """Ti si pomocni sistem u advokatskoj kancelariji koji klasifikuje dokumenta.
Dobices tekst izvucen iz dokumenta (moguce da OCR nije savrsen) i listu postojecih
klijenata kancelarije (format foldera: "Prezime Ime").

Zadatak:
1. Odredi kratak tip dokumenta na srpskom, prvo slovo veliko (npr. "Tuzba", "Presuda",
   "Resenje", "Ugovor", "Punomocje", "Zalba", "Predlog za izvrsenje",
   "Opomena pred izvrsenje", "Racun", "Zapisnik").
2. Odredi ime i prezime FIZICKOG LICA - klijenta kancelarije na koga se dokument
   odnosi (obicno stranka koju kancelarija zastupa, ne protivna strana, ne sud).
   Ako prepoznajes da je to neko sa liste postojecih klijenata, iskoristi TACNO
   to ime i prezime sa liste (isto pisanje).
3. Proceni sigurnost: "visoka" (jasno i nedvosmisleno), "srednja" (verovatno tacno
   ali postoji nesigurnost), "niska" (nejasno, vise mogucih lica, ili se ne moze
   pouzdano odrediti ime ili tip dokumenta).
4. Ako je sigurnost "niska" ili "srednja", u polju "napomena" kratko objasni zasto.

Odgovori ISKLJUCIVO validnim JSON-om, bez ikakvog dodatnog teksta, tacno ovog oblika:
{"tip_dokumenta": "...", "ime": "...", "prezime": "...", "sigurnost": "visoka|srednja|niska", "napomena": "..."}
"""


def log_append(log_putanja: str, linija: str) -> None:
    sada = datetime.now().strftime("%Y-%m-%d %H:%M")
    with open(log_putanja, "a", encoding="utf-8") as f:
        f.write(f"{sada} | {linija}\n")


# ─── Normalizacija imena ─────────────────────────────────────────────────────

def bez_dijakritika(s: str) -> str:
    s = s.replace("đ", "dj").replace("Đ", "Dj")
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


def normalizuj_za_poredjenje(s: str) -> str:
    s = bez_dijakritika(s).lower()
    return " ".join(s.split())


def sanitize_naziv(s: str) -> str:
    s = "".join(c for c in s if c not in ZABRANJENI_ZNACI_FAJL)
    return " ".join(s.split()).strip(" .")


# ─── Izvlacenje teksta iz dokumenata ─────────────────────────────────────────

def ocr_slika(pil_image) -> str:
    import pytesseract
    try:
        return pytesseract.image_to_string(pil_image, lang="srp+eng")
    except Exception:
        return pytesseract.image_to_string(pil_image, lang="eng")


def izvuci_tekst_pdf(putanja: str) -> str:
    import fitz  # PyMuPDF
    tekst_delovi = []
    doc = fitz.open(putanja)
    try:
        for page in doc:
            tekst_delovi.append(page.get_text())
        tekst = "\n".join(tekst_delovi).strip()
        if len(tekst) >= MIN_TEKST_ZNAKOVA:
            return tekst

        # PDF bez teksta (skeniran) - OCR po stranicama
        from PIL import Image
        ocr_delovi = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ocr_delovi.append(ocr_slika(img))
        return "\n".join(ocr_delovi).strip()
    finally:
        doc.close()


def izvuci_tekst_docx(putanja: str) -> str:
    import docx
    d = docx.Document(putanja)
    delovi = [p.text for p in d.paragraphs]
    for tabela in d.tables:
        for red in tabela.rows:
            delovi.append(" | ".join(c.text for c in red.cells))
    return "\n".join(delovi).strip()


def izvuci_tekst_slike(putanja: str) -> str:
    from PIL import Image
    with Image.open(putanja) as img:
        return ocr_slika(img).strip()


def izvuci_tekst(putanja: str):
    """Vraca (tekst, greska). Ako greska nije None, tekst se ignorise."""
    ext = os.path.splitext(putanja)[1].lower()
    try:
        if ext == ".pdf":
            return izvuci_tekst_pdf(putanja), None
        if ext == ".docx":
            return izvuci_tekst_docx(putanja), None
        if ext == ".doc":
            return None, "Stari .doc format nije podrzan - sacuvaj kao .docx ili PDF."
        if ext in (".jpg", ".jpeg", ".png"):
            return izvuci_tekst_slike(putanja), None
        return None, f"Nepodrzan format: {ext}"
    except Exception as e:
        return None, f"Greska pri citanju fajla: {e}"


# ─── Ollama ───────────────────────────────────────────────────────────────────

def pozovi_ollama(tekst: str, postojeci_klijenti: list, model: str) -> dict:
    poruka_korisnika = (
        "POSTOJECI KLIJENTI KANCELARIJE (folderi, format 'Prezime Ime'):\n"
        + ("\n".join(postojeci_klijenti) if postojeci_klijenti else "(nema jos nijednog)")
        + "\n\nTEKST DOKUMENTA:\n" + tekst[:MAX_TEKST_ZNAKOVA]
    )
    body = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": SISTEM_PROMPT},
            {"role": "user", "content": poruka_korisnika},
        ],
        "stream": False,
        "format": "json",
    }).encode("utf-8")

    req = urllib.request.Request(
        f"{OLLAMA_HOST}/api/chat", data=body,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    with urllib.request.urlopen(req, timeout=180) as resp:
        odgovor = json.loads(resp.read().decode("utf-8"))

    sadrzaj = odgovor["message"]["content"].strip()
    if sadrzaj.startswith("```"):
        sadrzaj = sadrzaj.strip("`")
        if sadrzaj.lower().startswith("json"):
            sadrzaj = sadrzaj[4:]
    return json.loads(sadrzaj.strip())


# ─── Poredjenje sa postojecim klijentima ─────────────────────────────────────

def ucitaj_postojece_klijente(klijenti_root: str) -> list:
    if not os.path.isdir(klijenti_root):
        return []
    return sorted(
        ime for ime in os.listdir(klijenti_root)
        if os.path.isdir(os.path.join(klijenti_root, ime)) and not ime.startswith("_")
    )


def resolvuj_klijenta(prezime: str, ime: str, postojeci: list):
    """Vraca (status, naziv_foldera). status: 'postojeci' | 'nejasno' | 'novi'."""
    ciljni_naziv = sanitize_naziv(f"{bez_dijakritika(prezime).strip()} {bez_dijakritika(ime).strip()}")
    cilj_norm = normalizuj_za_poredjenje(ciljni_naziv)

    najbolji, najbolji_ratio = None, 0.0
    for folder in postojeci:
        ratio = SequenceMatcher(None, cilj_norm, normalizuj_za_poredjenje(folder)).ratio()
        if ratio > najbolji_ratio:
            najbolji, najbolji_ratio = folder, ratio

    if najbolji and najbolji_ratio >= PODUDARANJE_POSTOJECI:
        return "postojeci", najbolji
    if najbolji and najbolji_ratio >= PODUDARANJE_SLICNO:
        return "nejasno", najbolji
    return "novi", ciljni_naziv


# ─── Obrada jednog fajla ──────────────────────────────────────────────────────

def jedinstvena_putanja(putanja: str) -> str:
    if not os.path.exists(putanja):
        return putanja
    koren, ext = os.path.splitext(putanja)
    i = 2
    while os.path.exists(f"{koren} ({i}){ext}"):
        i += 1
    return f"{koren} ({i}){ext}"


def premesti_u_sporno(putanja: str, sporno_dir: str, razlog: str, log: str) -> None:
    os.makedirs(sporno_dir, exist_ok=True)
    ciljna = jedinstvena_putanja(os.path.join(sporno_dir, os.path.basename(putanja)))
    shutil.move(putanja, ciljna)
    napomena_putanja = os.path.splitext(ciljna)[0] + ".napomena.txt"
    with open(napomena_putanja, "w", encoding="utf-8") as f:
        f.write(f"Fajl: {os.path.basename(putanja)}\nRazlog: {razlog}\n")
    linija = f"SPORNO: {os.path.basename(putanja)} | {razlog}"
    print(f"  [SPORNO] {razlog}")
    log_append(log, linija)


def obradi_fajl(putanja: str, klijenti_root: str, sporno_dir: str, model: str, log: str) -> None:
    naziv = os.path.basename(putanja)
    print(f"Obradjujem: {naziv}")

    tekst, greska = izvuci_tekst(putanja)
    if greska:
        premesti_u_sporno(putanja, sporno_dir, greska, log)
        return
    if not tekst or len(tekst.strip()) < MIN_TEKST_ZNAKOVA:
        premesti_u_sporno(putanja, sporno_dir, "Nije moguce izvuci citljiv tekst iz dokumenta.", log)
        return

    postojeci = ucitaj_postojece_klijente(klijenti_root)
    try:
        rezultat = pozovi_ollama(tekst, postojeci, model)
    except (urllib.error.URLError, TimeoutError) as e:
        print(f"  [GRESKA] Ollama nedostupna ({e}) - fajl ostaje na Desktopu, pokusace se ponovo.")
        log_append(log, f"GRESKA (Ollama nedostupna): {naziv} - {e}")
        return
    except Exception as e:
        print(f"  [GRESKA] Neuspesno prepoznavanje ({e}) - fajl ostaje na Desktopu.")
        log_append(log, f"GRESKA (parsiranje odgovora): {naziv} - {e}")
        return

    tip = (rezultat.get("tip_dokumenta") or "").strip()
    ime = (rezultat.get("ime") or "").strip()
    prezime = (rezultat.get("prezime") or "").strip()
    sigurnost = (rezultat.get("sigurnost") or "").strip().lower()
    napomena = (rezultat.get("napomena") or "").strip()

    if not tip or not ime or not prezime:
        premesti_u_sporno(putanja, sporno_dir, napomena or "AI nije uspeo da odredi tip dokumenta i/ili klijenta.", log)
        return
    if sigurnost == "niska":
        premesti_u_sporno(putanja, sporno_dir, napomena or "AI nije siguran u prepoznavanje (niska sigurnost).", log)
        return

    status, ciljni_folder = resolvuj_klijenta(prezime, ime, postojeci)
    if status == "nejasno":
        razlog = f"Ime '{prezime} {ime}' je slicno postojecem folderu '{ciljni_folder}' ali se ne poklapa tacno - proveri da nije isti klijent."
        premesti_u_sporno(putanja, sporno_dir, razlog, log)
        return

    folder_putanja = os.path.join(klijenti_root, ciljni_folder)
    os.makedirs(folder_putanja, exist_ok=True)

    ext = os.path.splitext(putanja)[1].lower()
    tip_sanitizovan = sanitize_naziv(tip)
    novi_naziv = f"{tip_sanitizovan} - {ciljni_folder} - {datetime.now().strftime('%d.%m.%Y')}{ext}"
    ciljna_putanja = jedinstvena_putanja(os.path.join(folder_putanja, novi_naziv))

    shutil.move(putanja, ciljna_putanja)

    novo_stanje = "nov folder" if status == "novi" else "postojeci folder"
    linija = f"{naziv} -> {ciljni_folder}/{os.path.basename(ciljna_putanja)} ({novo_stanje}, sigurnost: {sigurnost})"
    print(f"  -> {ciljni_folder}\\{os.path.basename(ciljna_putanja)}  ({novo_stanje})")
    log_append(log, linija)


# ─── Skeniranje Desktopa ──────────────────────────────────────────────────────

def treba_preskociti(entry: os.DirEntry) -> bool:
    if not entry.is_file():
        return True
    ext = os.path.splitext(entry.name)[1].lower()
    if ext not in DOZVOLJENE_EKSTENZIJE:
        return True
    if entry.name.startswith("~$"):  # Word lock fajl
        return True
    try:
        starost = datetime.now().timestamp() - entry.stat().st_mtime
    except OSError:
        return True
    if starost < MIN_STAROST_SEK:
        return True
    return False


def main():
    parser = argparse.ArgumentParser(description="Dokument agent - sortiranje dokumenata sa Desktopa")
    parser.add_argument("--desktop", default=DESKTOP_DIR_DEFAULT, help="Putanja do Desktopa")
    parser.add_argument("--klijenti", default=None, help="Putanja do foldera sa klijentima (default: <desktop>\\Advokat Zlatanovic)")
    parser.add_argument("--sporno", default=None, help="Putanja do foldera za sporne dokumente (default: <desktop>\\_Za proveru)")
    parser.add_argument("--model", default=os.environ.get("OLLAMA_MODEL", OLLAMA_MODEL_DEFAULT), help="Ollama model")
    parser.add_argument("--log", default=os.path.join(BASE_DIR, "dokument_log.txt"), help="Putanja log fajla")
    args = parser.parse_args()

    klijenti_root = args.klijenti or os.path.join(args.desktop, KLIJENTI_FOLDER_NAZIV)
    sporno_dir = args.sporno or os.path.join(args.desktop, SPORNO_FOLDER_NAZIV)

    print("=" * 70)
    print(" Dokument agent - sortiranje dokumenata")
    print("=" * 70)

    if not os.path.isdir(args.desktop):
        print(f"[GRESKA] Desktop folder ne postoji: {args.desktop}")
        sys.exit(1)

    try:
        fajlovi = [e for e in os.scandir(args.desktop) if not treba_preskociti(e)]
    except OSError as e:
        print(f"[GRESKA] Ne mogu da citam Desktop: {e}")
        sys.exit(1)

    if not fajlovi:
        print("Nema novih dokumenata na Desktopu.")
        return

    print(f"Pronadjeno {len(fajlovi)} novih dokument(a).\n")
    for entry in fajlovi:
        obradi_fajl(entry.path, klijenti_root, sporno_dir, args.model, args.log)

    print("\nGotovo.")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n[GRESKA] {e}")
        import traceback
        traceback.print_exc()
        try:
            log_append(os.path.join(BASE_DIR, "dokument_log.txt"), f"GRESKA: {e}")
        except Exception:
            pass
        sys.exit(1)
